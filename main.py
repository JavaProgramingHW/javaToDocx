# 이미지 추가 :
# from docx.shared import Inches
# document.add_picture('path', width=Inches(1.25))

import os
import re
import time
import json
import shutil
import logging
import datetime
import multiprocessing
from docx import Document
from docx.shared import RGBColor, Pt, Cm

from utils.get_file_list import get_file_list
from utils.get_command import get_command
from utils.get_java_response import get_java_response, get_java_pk_response
from utils.uml import *

version = "2.0.1"

class Java2Docx:
    def __init__(self):
        self.LOGGER = None
        self.userdata_path = "userdata.json"
        self.temp_folder_for_uml = "temp"
        self.uml_img_path = "umlimg"
        self.java_response_path = "console_img"

        self.path = None
        self.targets = []

        self.document = None

        self.homework_name = None
        self.grade = None
        self.studentID = None
        self.name = None
    
    def setup(self):
        # userdata 파일이 있을때
        if os.path.exists(self.userdata_path):
            file = open(self.userdata_path, "r")
            data = file.read()
            file.close()
            json_object = json.loads(data)
            if json_object["year"] == str(datetime.datetime.now().year):
                self.grade = json_object["grade"]
                self.studentID = json_object["studentID"]
                self.name = json_object["name"]

        # userdata 파일이 없을때
        if self.name is None:
            self.grade = input("학년을 입력하세요(ex. 1) : ").strip()
            self.studentID = input("학번을 입력하세요(ex. 20220101) : ").strip()
            self.name = input("이름을 입력하세요(ex. 홍길동) : ").strip()

            file = open(self.userdata_path, "w")
            file.write(f"""{{
            "year": "{datetime.datetime.today().year}",
            "grade": "{self.grade}",
            "studentID": "{self.studentID}",
            "name": "{self.name}"
        }}
        """)
            file.close()
        
        self.targets = get_file_list(self.path)
        self.LOGGER.info(f"대상 파일 리스트 - {self.targets}")

        # 도큐먼트 선언
        self.document = Document()

        p = self.document.add_heading(level=0)
        wp = p.add_run("자바프로그래밍")
        wp.font.color.rgb = RGBColor(0, 0, 0)
        p = self.document.add_heading(level=1)
        wp = p.add_run(self.homework_name)
        wp.font.color.rgb = RGBColor(0, 0, 0)
        
        p = self.document.add_heading(level=3)
        wp = p.add_run(f"금오공과대학교 컴퓨터소프트웨어공학과\n{self.grade}학년 {self.studentID} {self.name}")
        wp.font.color.rgb = RGBColor(0, 0, 0)


    def make_uml_class_diagram(self):
        # UML class diagram 이미지 저장할 임시폴더 생성
        try:
            shutil.rmtree(self.uml_img_path)
        except FileNotFoundError:
            pass
        os.mkdir(self.uml_img_path)

        # 코드 저장할 임시폴더 생성(UML class diagram 작성을 위함)
        try:
            shutil.rmtree(self.temp_folder_for_uml)
        except FileNotFoundError:
            pass
        os.mkdir(self.temp_folder_for_uml)

        process_list = []

        for file in self.targets:
            # 패키지일 경우
            if not file.endswith(".java"):
                # 폴더 째로 복사
                shutil.copytree(f"{self.path}/{file}", f"{self.temp_folder_for_uml}/{file}")

                # 임시로 만든 폴더에서 실행
                from_file = f"{self.temp_folder_for_uml}/{file}"
                # UML class diagram 이미지 저장할 폴더에 저장
                to_img = f"{self.uml_img_path}/{file}.png"
            # 단일 파일일 경우
            else:
                # 폴더명 생성
                temp_file_name = file.replace(".java", "")
                # 폴더 생성
                if not os.path.exists(f"{self.temp_folder_for_uml}/{temp_file_name}"):
                    os.mkdir(f"{self.temp_folder_for_uml}/{temp_file_name}")
                shutil.copy(f"{self.path}/{file}", f"{self.temp_folder_for_uml}/{temp_file_name}/{file}")

                # 임시로 만든 폴더에서 실행
                from_file = f"{self.temp_folder_for_uml}/{temp_file_name}"
                # UML class diagram 이미지 저장할 폴더에 저장
                to_img = f"{self.uml_img_path}/{temp_file_name}.png"
            
            # 이미지 다운로드 프로세스 생성
            process = multiprocessing.Process(target=download_uml_class_diagram_img, args=(from_file, to_img))
            process.start()
            process_list.append(process)
        
        return process_list
    
    def make_java_response(self):
        # 코드 결과값 저장할 폴더 생성
        try:
            shutil.rmtree(self.java_response_path)
        except FileNotFoundError:
            pass
        os.mkdir(self.java_response_path)

        process_list = []

        for file in self.targets:
            # 패키지일 경우
            if not file.endswith(".java"):
                if os.path.exists(f"{self.path}/{file}/{file}.java"):
                    # 커맨드 가져오기 위해 코드 불러오기
                    code_date = open(f"{self.path}/{file}/{file}.java", "r")
                    code = code_date.read()
                    code_date.close()
                    
                    # 코드에서 커맨드 가져오기
                    _, command = get_command(code)

                    process = multiprocessing.Process(target=get_java_pk_response, args=(self.path, file, command, self.java_response_path))
                    process.start()
                    process_list.append(process)
            
            # 단일 파일일 경우
            else:
                if os.path.exists(f"{self.path}/{file}"):
                    # 커맨드 가져오기 위해 코드 불러오기
                    code_date = open(f"{self.path}/{file}", "r")
                    code = code_date.read()
                    code_date.close()

                    # 코드에서 커맨드 가져오기
                    _, command = get_command(code)

                    process = multiprocessing.Process(target=get_java_response, args=(self.path, file, command, self.java_response_path))
                    process.start()
                    process_list.append(process)
        
        return process_list
    
    def make_document(self):
        for file in self.targets:
            temp_path = None
            temp_file_list = None

            self.document.add_page_break() # 다음 페이지로

            # 문제 제목 뽑아내기
            temp = file.split("_")
            after_list = []
            for i in temp:
                after_list.append(re.sub(r'[^0-9]', '', i))
            cell_title = ".".join(after_list)

            # 문제 제목 넣기 - ex. 1.1
            p = self.document.add_heading(level=1)
            wp = p.add_run(cell_title)
            wp.font.size = Pt(20) # 글자 크기 조절
            wp.font.color.rgb = RGBColor(0, 0, 0) # 글자 색깔 검정색으로

            # 파일이 패키지일 경우
            if not file.endswith(".java"):
                # 경로 설정
                temp_path = f"{self.path}/{file}"
                temp_file_list = get_file_list(temp_path)
            
                temp_file_list_processing = []
                file_java_exist = False
                for i in temp_file_list:
                    # 만약 폴더명과 동일한 이름의 java 파일이 아닐 경우 추가 - 이를 리스트의 맨 위에 올리기 위함
                    if i != f"{file}.java":
                        temp_file_list_processing.append(i)
                    else:
                        file_java_exist = True

                # 리스트의 맨 앞에 추가
                if file_java_exist:
                    temp_file_list_processing.insert(0, f"{file}.java")

            # 테이블 작성
            table = self.document.add_table(rows=1, cols=1)
            table.style = "Table Grid"

            # 테이블에 UML class diagram 이미지 추가
            hdr_cells = table.rows[0].cells

            if os.path.exists(f"{self.uml_img_path}/{file}.png".replace(".java", "")):
                paragraph = hdr_cells[0].paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(f"{self.uml_img_path}/{file}.png".replace(".java", ""), width=Cm(15))
            else:
                hdr_cells[0].text = f"설계 : UML class diagram - 이미지 생성 실패"

            # 테이블에 코드 추가
            # 패키지일 경우
            if temp_path is not None:
                for i in enumerate(temp_file_list_processing):
                    # 코드 가져오기
                    code_date = open(f"{temp_path}/{i[1]}", "r")
                    code = code_date.read()
                    code_date.close()

                    # 커맨드 적은 주석 제거
                    if i[1] == f"{file}.java":
                        end_index, _ = get_command(code)

                    if end_index is not None:
                        temp_code = code.split("\n")
                        code = "\n".join(temp_code[end_index+1:])

                    row_cells = table.add_row().cells
                    row_cells[0].text = f"//{i[1]}\n\n{code}"
            else:
                # 코드 가져오기
                code_date = open(f"{self.path}/{file}", "r")
                code = code_date.read()
                code_date.close()

                # 커맨드 적은 주석 제거
                end_index, _ = get_command(code)

                if end_index is not None:
                    temp_code = code.split("\n")
                    code = "\n".join(temp_code[end_index + 1:])
                
                row_cells = table.add_row().cells
                row_cells[0].text = code
            
            # 결과값
            row_cells = table.add_row().cells
            image_path = f"{self.java_response_path}/{file}.png".replace(".java", "")
            if os.path.exists(image_path):
                paragraph = row_cells[0].paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(image_path, width=Cm(15))
            else:
                row_cells[0].text = f"실행 결과 : 이미지 생성 실패"
            
            row_cells = table.add_row().cells
            row_cells[0].text = f"난이도 : 중"

            row_cells = table.add_row().cells
            row_cells[0].text = f"완성도 : 정상 실행됨"

        self.document.save('result.docx')
    
    def finish_clear(self):
        # 임시폴더 제거
        if os.path.exists(self.temp_folder_for_uml):
            shutil.rmtree(self.temp_folder_for_uml)

        if os.path.exists(self.uml_img_path):
            shutil.rmtree(self.uml_img_path)

        if os.path.exists(self.java_response_path):
            shutil.rmtree(self.java_response_path)

if __name__ == "__main__":
    main_class = Java2Docx()

    # enable logging
    logging.basicConfig(
        format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
        level=logging.INFO)
    main_class.LOGGER = logging.getLogger(__name__)

    os.system("java --version")
    print()
    main_class.LOGGER.info(f"javaToDocx Version {version}")
    main_class.LOGGER.info("Made by @ajb3296")
    print()
    main_class.path = input("대상 경로를 입력하세요 : ").strip()
    main_class.homework_name = input("과제명을 입력하세요 : ").strip()

    start = time.time()

    main_class.setup()

    # UML class diagram 이미지 생성 프로세스 실행
    uml_process_list = main_class.make_uml_class_diagram()

    # 자바 실행 결과 이미지 생성 프로세스 실행
    java_response_process_list = main_class.make_java_response()

    # 프로세스 기다림
    for process in uml_process_list:
        process.join()
    for process in java_response_process_list:
        process.join()

    main_class.make_document()

    main_class.finish_clear()

    main_class.LOGGER.info("완료")
    main_class.LOGGER.info(f"소요시간 : {time.time() - start} 초")